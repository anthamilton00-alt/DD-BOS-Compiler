from pathlib import Path
import re

from docx import Document

from core.registry import (
    DATE_LABELS,
    DOCUMENT_ID_LABELS,
    INVALID_OWNER_VALUES,
    OWNER_LABELS,
    STATUS_LABELS,
    TITLE_LABELS,
    VALID_DOCUMENT_PREFIXES,
    VALID_OWNERS,
    VALID_STATUSES,
    VERSION_LABELS,
)


METADATA_FIELDS = {
    "Document ID": "",
    "Title": "",
    "Version": "",
    "Status": "",
    "Owner": "",
    "Effective Date": "",
    "Revision Date": "",
}


LABEL_MAP = {
    **{label: "Document ID" for label in DOCUMENT_ID_LABELS},
    **{label: "Title" for label in TITLE_LABELS},
    **{label: "Version" for label in VERSION_LABELS},
    **{label: "Status" for label in STATUS_LABELS},
    **{label: "Owner" for label in OWNER_LABELS},
    "EFFECTIVE DATE": "Effective Date",
    "DATE EFFECTIVE": "Effective Date",
    "REVISION DATE": "Revision Date",
    "LAST REVISION DATE": "Revision Date",
    "LAST REVISED": "Revision Date",
}


PREFIX_PATTERN = "|".join(
    sorted(
        (re.escape(prefix) for prefix in VALID_DOCUMENT_PREFIXES),
        key=len,
        reverse=True,
    )
)

DOCUMENT_ID_PATTERN = re.compile(
    rf"\b(?:{PREFIX_PATTERN})-\d+\b",
    re.IGNORECASE,
)

VERSION_PATTERN = re.compile(
    r"^[Vv]?\d+(?:\.\d+)*$"
)

DATE_PATTERN = re.compile(
    r"^(?:"
    r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|"
    r"\d{4}-\d{1,2}-\d{1,2}"
    r"|"
    r"[A-Za-z]+\s+\d{1,2},\s+\d{4}"
    r"|"
    r"\d{1,2}\s+[A-Za-z]+\s+\d{4}"
    r")$"
)


INVALID_VALUES = {
    "",
    "-",
    "--",
    "---",
    "N/A",
    "NA",
    "NONE",
    "TBD",
    "TO BE DETERMINED",
    "DATE",
    "DATE APPROVED",
    "APPROVED BY",
    "PREPARED BY",
    "DESCRIPTION",
    "PURPOSE",
    "SCOPE",
}


VALID_OWNER_LOOKUP = {
    owner.upper(): owner
    for owner in VALID_OWNERS
}

INVALID_OWNER_LOOKUP = {
    owner.upper()
    for owner in INVALID_OWNER_VALUES
}


def normalize_text(value: str) -> str:
    """
    Normalize whitespace without removing metadata labels.
    """

    return " ".join(value.split()).strip()


def normalize_label(value: str) -> str:
    """
    Normalize a metadata label for matching.
    """

    value = normalize_text(value)
    value = value.rstrip(":").strip()

    return value.upper()


def clean_value(value: str) -> str:
    """
    Normalize and validate a possible metadata value.
    """

    value = normalize_text(value)
    value = value.strip("|").strip()

    if normalize_label(value) in INVALID_VALUES:
        return ""

    if value and all(character in {"_", "-", " "} for character in value):
        return ""

    return value


def normalize_owner(value: str) -> str:
    """
    Return the registry-approved owner name or an empty string.
    """

    value = clean_value(value)

    if not value:
        return ""

    normalized = normalize_label(value)

    if normalized in INVALID_OWNER_LOOKUP:
        return ""

    return VALID_OWNER_LOOKUP.get(normalized, "")


def extract_filename_metadata(doc_path):
    """
    Extract authoritative ID and title values from the filename.
    """

    stem = Path(doc_path).stem.strip()

    document_id = ""
    title = stem

    match = DOCUMENT_ID_PATTERN.search(stem)

    if match:
        document_id = match.group(0).upper()

        title = re.sub(
            rf"^{re.escape(match.group(0))}\s*[-—–]?\s*",
            "",
            stem,
            flags=re.IGNORECASE,
        ).strip()

    return document_id, title


def collect_metadata_lines(doc) -> list[str]:
    """
    Collect paragraph and table-cell text while preserving labels.
    """

    lines: list[str] = []

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)

        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = normalize_text(paragraph.text)

                    if text:
                        lines.append(text)

    return lines


def assign_metadata_value(
    metadata: dict[str, str],
    field: str,
    value: str,
) -> None:
    """
    Validate and assign a metadata field.
    """

    value = clean_value(value)

    if not value:
        return

    if normalize_label(value) in LABEL_MAP:
        return

    if field == "Document ID":
        if metadata["Document ID"]:
            return

        match = DOCUMENT_ID_PATTERN.search(value)

        if match:
            metadata["Document ID"] = match.group(0).upper()

        return

    if field == "Title":
        if not metadata["Title"]:
            metadata["Title"] = value

        return

    if field == "Version":
        normalized = value.lstrip("Vv")

        if VERSION_PATTERN.fullmatch(value):
            metadata["Version"] = normalized

        return

    if field == "Status":
        normalized = value.upper()

        if normalized in VALID_STATUSES:
            metadata["Status"] = normalized

        return

    if field in {"Effective Date", "Revision Date"}:
        if DATE_PATTERN.fullmatch(value):
            metadata[field] = value

        return

    if field == "Owner":
        if metadata["Owner"]:
            return

        owner = normalize_owner(value)

        if owner:
            metadata["Owner"] = owner


def extract_metadata(doc_path):
    """
    Extract DD-BOS metadata from paragraphs and tables.

    The filename remains authoritative for Document ID and Title.
    Registry values control document prefixes, statuses, labels, and owners.
    """

    doc = Document(doc_path)
    metadata = METADATA_FIELDS.copy()

    filename_id, filename_title = extract_filename_metadata(doc_path)

    metadata["Document ID"] = filename_id
    metadata["Title"] = filename_title

    lines = collect_metadata_lines(doc)

    for index, text in enumerate(lines):
        normalized_text = normalize_label(text)

        if normalized_text in LABEL_MAP:
            if index + 1 < len(lines):
                assign_metadata_value(
                    metadata,
                    LABEL_MAP[normalized_text],
                    lines[index + 1],
                )

            continue

        inline_match = re.match(
            r"^\s*([^:]{2,40})\s*:\s*(.+?)\s*$",
            text,
        )

        if not inline_match:
            continue

        label = normalize_label(inline_match.group(1))
        value = inline_match.group(2)

        if label in LABEL_MAP:
            assign_metadata_value(
                metadata,
                LABEL_MAP[label],
                value,
            )

    return metadata