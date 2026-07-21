from pathlib import Path
import re

from docx import Document

from core.config import load_config
from core.metadata import extract_metadata
from core.models import DocumentRecord


REFERENCE_PATTERN = re.compile(
    r"\b(?:ASM|ROOM|DD|FP|LAB|PROD)-\d+\b",
    re.IGNORECASE,
)


def is_section_heading(text: str) -> bool:
    """
    DD-BOS section headings are typically uppercase lines.
    """

    text = text.strip()

    if len(text) < 3:
        return False

    if ":" in text:
        return False

    letters = [c for c in text if c.isalpha()]

    if not letters:
        return False

    return all(c.isupper() for c in letters)


def scan_documents():
    """
    Scan all DD-BOS documents.
    """

    config = load_config()
    docs_folder = config.input_folder

    records = []

    print("\nScanning documents...\n")

    if not docs_folder.exists():
        print(f"Document folder not found:\n{docs_folder}\n")
        return records

    for file in sorted(docs_folder.glob("*.docx")):

        try:

            doc = Document(file)
            metadata = extract_metadata(file)

            sections = []
            references = set()

            for paragraph in doc.paragraphs:

                text = paragraph.text.strip()

                if not text:
                    continue

                if is_section_heading(text):
                    sections.append(text)

                for match in REFERENCE_PATTERN.findall(text):
                    references.add(match.upper())

            record = DocumentRecord(
                file_path=file,
                file_name=file.name,
                document_id=metadata["Document ID"] or file.stem.split(" ")[0],
                document_type=file.stem.split("-")[0],
                title=metadata["Title"],
                version=metadata["Version"],
                status=metadata["Status"] or "OK",
                owner=metadata["Owner"],
                effective_date=metadata["Effective Date"],
                revision_date=metadata["Revision Date"],
                sections=sections,
                references=sorted(references),
                paragraph_count=len(doc.paragraphs),
                table_count=len(doc.tables),
            )

            records.append(record)

            print(f"✓ {record.file_name}")

        except Exception as ex:

            record = DocumentRecord(
                file_path=file,
                file_name=file.name,
            )

            record.errors.append(str(ex))

            records.append(record)

            print(f"✗ {file.name}")
            print(f"    {ex}")

    print("\nFinished.\n")
    print(f"{len(records)} documents processed.\n")

    return records